import os
import shutil
import datetime

from django.conf import settings

import login.files_upload as files_upload
from docx import Document
import re


def create_job_folders(master_folder, job_number, job_title):
    job_folder_name = f"{job_number} {job_title}"
    job_folder_path = os.path.join(master_folder, job_folder_name)

    if os.path.exists(job_folder_path):
        shutil.rmtree(job_folder_path)
        print(f"Existing job folder '{job_folder_path}' removed.")

    os.makedirs(job_folder_path)

    subfolders = ['Quote', 'Visual', 'RAMS', 'Client Documents', 'PO']
    for subfolder in subfolders:
        subfolder_path = os.path.join(job_folder_path, subfolder)
        os.makedirs(subfolder_path)
        print(f"Created subfolder: {subfolder_path}")

    print(f"Created job folder: {job_folder_path}")


def replace_text_in_table(docx_path, replacements, job_number):
    doc = Document(docx_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for pattern, replacement in replacements.items():
                    new_text = re.sub(pattern, replacement, cell.text)
                    cell.text = new_text
    for section in doc.sections:
        footer = section.footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'Texttexttext4' in cell.text:
                        cell.text = cell.text.replace('Texttexttext4', str(datetime.date.today()))
                    elif 'Texttexttext5' in cell.text:
                        cell.text = cell.text.replace('Texttexttext5', job_number)

    return doc


def main(jobNumber, jobDescription, address, dateOfWorks, duration, localHospital, natureOfWorks, materials):
    master_folder = 'upload_files'
    job_number = jobNumber
    job_title = jobDescription
    folder_name = f"{job_number} {job_title}"
    create_job_folders(master_folder, job_number, job_title)
    # document_path1 = os.path.join(settings.MEDIA_URL, 'document1.docx')
    # document_path2 = os.path.join(settings.MEDIA_URL, 'document2.docx')

    input_docx_path1 = 'static/docs/document1.docx'
    output_docx_path1 = f'upload_files/{folder_name}/RAMS/{folder_name} - RA.docx'
    input_docx_path2 = 'static/docs/document2.docx'
    output_docx_path2 = f'upload_files/{folder_name}/RAMS/{folder_name} - MS.docx'

    job_number = jobNumber

    pattern_replacements = {
        r'\bTexttexttext1\b': folder_name,
        r'\bTexttexttext2\b': dateOfWorks,
        r'\bTexttexttext3\b': duration,
        r'\bTexttexttext4\b': str(datetime.date.today()),
        r'\bTexttexttext5\b': jobNumber,
        r'\bTexttexttext6\b': localHospital,
        r'\bTexttexttext7\b': address,
        r'\bTexttexttext8\b': jobDescription,
        r'\bTexttexttext9\b': natureOfWorks,
        r'\bTexttexttext10\b': materials,
    }

    modified_docx = replace_text_in_table(input_docx_path1, pattern_replacements, job_number)
    modified_docx.save(output_docx_path1)
    modified_docx = replace_text_in_table(input_docx_path2, pattern_replacements, job_number)
    modified_docx.save(output_docx_path2)

    # RA Other
    ra_files = [
        'Flammable Liquids (use and storage of) -  RA.docx', 'Lone Working (out of office) - RA.docx', 'Manual Handling (pushing & pulling) - RA.docx',
        'Mobile Elevated Work Platform (use of) - RA.docx', 'Occupied Property (working in) - RA.docx',
        'Saws (working with) - RA.docx', 'Stress (work related) - RA.docx', 'Surveying Pack - RA.docx',
        'Vibration (hand arm) - RA.docx', 'Vinyl Plotting Machine (working with) - RA.docx',
        'Work at Height (scaffolding) - RA.docx', 'Young Persons - RA.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Risk Assesments/RA Other/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Risk Assesments/RA Other/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Risk Assesments", "Risk Assesments/RA Other"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # Additional Information
    ra_files = ['Risk Assessment - Additional Information.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Risk Assesments/Additional Information/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Risk Assesments/Additional Information/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Risk Assesments", "Risk Assesments/Additional Information"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # External Works
    ra_files = ['Architectural Film Installation (external) - MS.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Method Statements/External Works/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Method Statements/External Works/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Method Statements", "Method Statements/External Works"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # Internal Lift Specific
    ra_files = ['Lift refurbishment - MS.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Method Statements/Internal Lift Specific/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Method Statements/Internal Lift Specific/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Method Statements", "Method Statements/Internal Lift Specific"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # Internal Works
    ra_files = ['Architectural Film Installation (internal) - MS.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Method Statements/Internal Works/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Method Statements/Internal Works/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Method Statements", "Method Statements/Internal Works"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # Method Statement Checklist
    ra_files = ['MS Checklist.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Method Statements/Method Statement Checklist/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Method Statements/Method Statement Checklist/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Method Statements", "Method Statements/Method Statement Checklist"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # Surveying
    ra_files = ['Surveying - MS.docx']
    for file in ra_files:
        ra_file_path = f'static/docs/Archive/Method Statements/Surveying/{file}'
        output_path = f'upload_files/{folder_name}/RAMS/Method Statements/Surveying/{file}'
        base_path = f"upload_files/{folder_name}/RAMS"
        subfolders = ["Method Statements", "Method Statements/Surveying"]
        for subfolder in subfolders:
            folder_path = os.path.join(base_path, subfolder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
        modified_docx = replace_text_in_table(ra_file_path, pattern_replacements, job_number)
        modified_docx.save(output_path)

    # COSHH files
    source_directory = 'static/docs/Archive/COSHH Assessments'
    destination_directory = f'upload_files/{folder_name}/RAMS/COSHH'

    if not os.path.exists(destination_directory):
        try:
            shutil.copytree(source_directory, destination_directory)
            print(f"Directory '{source_directory}' copied to '{destination_directory}'")
        except OSError as e:
            print(f"Error copying directory: {e}")
    else:
        print(f"Destination directory '{destination_directory}' already exists.")

    files_upload.main()
    return True

# main('123', 'desdes', 'addadd', 'september 2023', '2 months', 'hoshoshos', 'testing1', 'testing2')
